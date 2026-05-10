"""
generate_report_docx.py
-----------------------
Converts REPORT.md into a formatted, submittable Word (.docx) document.
Run: python3 generate_report_docx.py
Output: CS444_IR_Project_Report.docx
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_horizontal_rule(doc):
    p  = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


def shade_cell(cell, hex_color="1F3864"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_border(table):
    """Add thin borders to every cell in the table."""
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"),   "single")
                border.set(qn("w:sz"),    "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "AAAAAA")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def render_inline(para, text):
    """Render **bold**, *italic*, and `code` within a paragraph."""
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
    parts   = re.split(pattern, text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
            r.font.name = "Times New Roman"
        elif part.startswith("*") and part.endswith("*"):
            r = para.add_run(part[1:-1])
            r.italic = True
            r.font.name = "Times New Roman"
        elif part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        else:
            r = para.add_run(part)
            r.font.name = "Times New Roman"
        r.font.size = Pt(12)


# ─── Main builder ─────────────────────────────────────────────────────────────

def build_docx():
    doc = Document()

    # ── Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Default body style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # ═══════════════════════════════════════════════════════
    # COVER / TITLE BLOCK
    # ═══════════════════════════════════════════════════════
    def centre(text, size=12, bold=False, space_after=6, color=None):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(text)
        r.bold      = bold
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor(*color)
        return p

    centre("COMSATS University Islamabad", 13, bold=False, space_after=3)
    centre("Department of Computer Science", 12, space_after=10)
    add_horizontal_rule(doc)

    centre("CS444 — Information Retrieval", 14, bold=True, space_after=4,
           color=(31, 56, 100))
    centre("Semester Project Report", 18, bold=True, space_after=4,
           color=(31, 56, 100))
    centre("Classifier Comparison: Naive Bayes, Rocchio & k-Nearest Neighbors",
           13, space_after=14)

    add_horizontal_rule(doc)

    # ── Info table
    info_tbl = doc.add_table(rows=4, cols=2)
    info_tbl.style = "Table Grid"
    info_data = [
        ("Course",          "CS444 — Information Retrieval"),
        ("Instructor",      "Dr. Zoya"),
        ("Submission Type", "Semester Project"),
        ("Total Marks",     "100"),
    ]
    for i, (k, v) in enumerate(info_data):
        lc, rc = info_tbl.rows[i].cells
        lc.text = k
        rc.text = v
        lc.paragraphs[0].runs[0].bold      = True
        lc.paragraphs[0].runs[0].font.name = "Times New Roman"
        rc.paragraphs[0].runs[0].font.name = "Times New Roman"
        for cell in (lc, rc):
            cell.paragraphs[0].runs[0].font.size = Pt(12)
    set_cell_border(info_tbl)
    doc.add_paragraph()

    # ── Group members table
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hdr.add_run("Group Members")
    r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(12)

    grp_tbl = doc.add_table(rows=4, cols=2)
    grp_tbl.style = "Table Grid"
    headers  = ["Name", "Registration No."]
    members  = [
        ("Saad Hanif Taj", "2022509"),
        ("Ahmed Ali",      "2022054"),
        ("Aiza Azeem",     "2022077"),
    ]
    # Header row
    for j, h in enumerate(headers):
        cell = grp_tbl.rows[0].cells[j]
        cell.text = h
        r = cell.paragraphs[0].runs[0]
        r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(12)
        shade_cell(cell, "1F3864")
        r.font.color.rgb = RGBColor(255, 255, 255)

    for i, (name, reg) in enumerate(members):
        row = grp_tbl.rows[i + 1]
        row.cells[0].text = name
        row.cells[1].text = reg
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.name = "Times New Roman"
            cell.paragraphs[0].runs[0].font.size = Pt(12)
    set_cell_border(grp_tbl)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # READ + PARSE REPORT.md
    # ═══════════════════════════════════════════════════════
    with open("REPORT.md", encoding="utf-8") as f:
        lines = f.readlines()

    # Skip the header block already rendered (lines up to first ---)
    start = 0
    dashes_seen = 0
    for i, ln in enumerate(lines):
        if ln.strip() == "---":
            dashes_seen += 1
            if dashes_seen == 3:      # after cover table + group table separators
                start = i + 1
                break

    in_code   = False
    code_buf  = []
    in_table  = False
    tbl_rows  = []
    skip_next = False

    def flush_table():
        nonlocal in_table, tbl_rows
        if not tbl_rows:
            in_table = False
            return
        # Remove separator row (---)
        data_rows = [r for r in tbl_rows if not re.match(r"^\|[-| :]+\|$", r)]
        if not data_rows:
            in_table = False
            tbl_rows = []
            return

        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            parsed.append(cells)

        n_cols = max(len(r) for r in parsed)
        table  = doc.add_table(rows=len(parsed), cols=n_cols)
        table.style = "Table Grid"

        for ri, row in enumerate(parsed):
            for ci, cell_txt in enumerate(row):
                if ci >= n_cols:
                    break
                cell = table.rows[ri].cells[ci]
                p    = cell.paragraphs[0]
                p.clear()
                # Strip bold markers
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_txt)
                run   = p.add_run(clean)
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
                if ri == 0:
                    run.bold = True
                    shade_cell(cell, "D9E1F2")

        set_cell_border(table)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        in_table = False
        tbl_rows = []

    def flush_code(buf):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(1)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        for ln in buf:
            run = p.add_run(ln.rstrip("\n"))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            p.add_run("\n")

    for raw in lines[start:]:
        line = raw.rstrip("\n")

        # Skip pure HR lines
        if line.strip() == "---":
            add_horizontal_rule(doc)
            continue

        # ── Code blocks
        if line.startswith("```"):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                flush_code(code_buf)
                in_code  = False
                code_buf = []
            continue
        if in_code:
            code_buf.append(line)
            continue

        # ── Tables
        if line.startswith("|"):
            if not in_table:
                in_table = True
                tbl_rows = []
            tbl_rows.append(line)
            continue
        else:
            if in_table:
                flush_table()

        # ── Blank line
        if not line.strip():
            continue

        # ── Headings
        if line.startswith("#### "):
            p = doc.add_heading(line[5:], level=4)
            p.runs[0].font.name = "Times New Roman"
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
            p.runs[0].font.name = "Times New Roman"
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            p.runs[0].font.name = "Times New Roman"
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.name = "Times New Roman"
            continue

        # ── Bullet points
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent  = Cm(0.5)
            p.paragraph_format.space_after  = Pt(2)
            render_inline(p, line[2:])
            continue

        # ── Numbered list
        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(2)
            render_inline(p, re.sub(r"^\d+\.\s", "", line))
            continue

        # ── Block quote / formula lines (start with spaces or $)
        if line.startswith("    ") or line.startswith("$$"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(line.strip())
            run.font.name  = "Courier New"
            run.font.size  = Pt(10)
            run.font.italic = True
            continue

        # ── Normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
        render_inline(p, line)

    # flush any remaining table
    if in_table:
        flush_table()

    # ─── Footer with page numbers
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()
        r1 = fp.add_run("CS444 — Information Retrieval  |  Group: Saad Hanif Taj, Ahmed Ali, Aiza Azeem  |  Page ")
        r1.font.name = "Times New Roman"; r1.font.size = Pt(9)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run_xml = fp.add_run()._r
        run_xml.append(fldChar1)
        run_xml.append(instrText)
        run_xml.append(fldChar2)

    out = "CS444_IR_Project_Report.docx"
    doc.save(out)
    print(f"✓  Saved: {out}")


if __name__ == "__main__":
    build_docx()
